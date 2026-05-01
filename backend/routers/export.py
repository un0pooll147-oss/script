from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import io

router = APIRouter()

class ExportRequest(BaseModel):
    script: str
    title: str = "脚本"
    format: str = "pdf"  # "pdf" or "docx"


def build_pdf(title: str, script: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        # 日本語フォントを登録（システムにあるものを優先）
        font_candidates = [
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSans"),
            ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSans"),
            ("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc", "Hiragino"),
            ("/Library/Fonts/Arial Unicode MS.ttf", "ArialUnicode"),
        ]
        font_name = "Helvetica"  # fallback
        for path, name in font_candidates:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    font_name = name
                    break
                except Exception:
                    continue

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=25*mm, rightMargin=25*mm,
            topMargin=25*mm, bottomMargin=25*mm
        )

        title_style = ParagraphStyle(
            "Title", fontName=font_name, fontSize=18,
            leading=28, spaceAfter=8*mm, alignment=1  # center
        )
        body_style = ParagraphStyle(
            "Body", fontName=font_name, fontSize=11,
            leading=20, spaceAfter=4*mm
        )
        scene_style = ParagraphStyle(
            "Scene", fontName=font_name, fontSize=11,
            leading=20, spaceAfter=2*mm,
            leftIndent=0, textColor=(0.3, 0.3, 0.3)
        )

        story = [Paragraph(title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), title_style)]
        story.append(Spacer(1, 4*mm))

        for line in script.split("\n"):
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if safe.strip() == "":
                story.append(Spacer(1, 3*mm))
            elif safe.startswith("【") or safe.startswith("■") or safe.startswith("◆"):
                story.append(Paragraph(f"<b>{safe}</b>", body_style))
            elif safe.startswith("　") or safe.startswith("  "):
                story.append(Paragraph(safe, scene_style))
            else:
                story.append(Paragraph(safe, body_style))

        doc.build(story)
        return buf.getvalue()

    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab が未インストールです: pip install reportlab")


def build_docx(title: str, script: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Mm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ページ余白
        for section in doc.sections:
            section.top_margin = Mm(25)
            section.bottom_margin = Mm(25)
            section.left_margin = Mm(25)
            section.right_margin = Mm(25)

        # タイトル
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        doc.add_paragraph()

        # 本文
        for line in script.split("\n"):
            if line.strip() == "":
                doc.add_paragraph()
                continue

            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(11)

            # 見出し行は太字
            if line.startswith("【") or line.startswith("■") or line.startswith("◆"):
                run.font.bold = True
            # シーン行はグレー
            elif line.startswith("　") or line.startswith("  "):
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx が未インストールです: pip install python-docx")


@router.post("/")
def export_script(req: ExportRequest):
    if not req.script.strip():
        raise HTTPException(status_code=400, detail="脚本が空です")

    safe_title = req.title.replace("/", "_").replace("\\", "_")[:50] or "script"

    if req.format == "pdf":
        data = build_pdf(req.title, req.script)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'}
        )
    elif req.format == "docx":
        data = build_docx(req.title, req.script)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'}
        )
    else:
        raise HTTPException(status_code=400, detail="format は 'pdf' または 'docx' を指定してください")
